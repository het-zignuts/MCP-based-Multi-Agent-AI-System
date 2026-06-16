import os
import tempfile
import pytest
from uuid import UUID

from docx import Document

from app.services.file_processing.file_processors import process_docx
from app.services.memory.memory_promoter import promote_memories_from_messages


def test_process_docx_handles_heading_styles_and_normal_paragraphs():
    with tempfile.TemporaryDirectory() as tmpdir:
        file_path = os.path.join(tmpdir, "test.docx")
        document = Document()
        heading = document.add_paragraph("Document Title")
        heading.style = document.styles["Heading 1"]
        normal = document.add_paragraph("This is a paragraph.")
        document.save(file_path)

        chunks = process_docx(file_path)

        joined = "\n".join(chunks)
        assert "# Document Title" in joined
        assert "This is a paragraph." in joined


class DummyMemoryItem:
    def __init__(self, data):
        self._data = data

    def model_dump(self):
        return self._data


class DummyCreatedMemory:
    def __init__(self, memory_id, content, memory_type, memory_metadata):
        self.id = memory_id
        self.content = content
        self.memory_type = memory_type
        self.memory_metadata = memory_metadata


@pytest.mark.asyncio
async def test_promote_memories_from_messages_accepts_dict_and_model_items(monkeypatch):
    dummy_messages = [type("Message", (), {"role": "user", "content": "I love pizza."})()]

    async def dummy_extract_memories_from_text(_text):
        return [
            DummyMemoryItem(
                {
                    "content": "I love pizza.",
                    "memory_type": "preference",
                    "importance_score": 0.7,
                    "confidence_score": 0.8,
                    "evidence": "explicit",
                    "temporal_scope": "durable",
                    "memory_metadata": {},
                }
            ),
            {
                "content": "I love pizza.",
                "memory_type": "preference",
                "importance_score": 0.7,
                "confidence_score": 0.8,
                "evidence": "explicit",
                "temporal_scope": "durable",
                "memory_metadata": {},
            },
        ]

    async def dummy_create_memory_with_embedding(
        db,
        user_id,
        conversation_id,
        content,
        memory_type,
        memory_metadata,
        importance_score,
        source,
        comparison_budget,
    ):
        return DummyCreatedMemory("123", content, memory_type, memory_metadata)

    monkeypatch.setattr(
        "app.services.memory.memory_promoter.extract_memories_from_text",
        dummy_extract_memories_from_text,
    )
    monkeypatch.setattr(
        "app.services.memory.memory_promoter.create_memory_with_embedding",
        dummy_create_memory_with_embedding,
    )

    created_memories = await promote_memories_from_messages(
        db=None,
        user_id=UUID("12345678-1234-5678-1234-567812345678"),
        messages=dummy_messages,
        conversation_id=UUID("87654321-4321-6789-4321-678987654321"),
    )

    assert len(created_memories) == 2
    assert created_memories[0]["content"] == "I love pizza."
    assert created_memories[1]["content"] == "I love pizza."
